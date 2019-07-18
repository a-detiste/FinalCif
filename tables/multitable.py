# -*- coding: utf-8 -*-
#
# This program was innitially created by Nils Trapp and extended by
# Daniel Kratzert
#
import itertools as it
import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document, table
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm

# compiled with "Py -3 -m PyInstaller multitable.spec --onefile"
from cif.file_reader import CifContainer
from tables.mtools import cif_keywords_list, isfloat, this_or_quest
from tables.symm import SymmetryElement

"""
#TODO: 
* create nice .docx template
* completeness to 0.83 A, theta range in A
* index ranges in one row
* cell parameters in three rows
* proper table numbering:
 table1 -> num -> table2 -> ...
 if table fails, return number, else count up and then return number 
"""


def populate_description_columns(table, cif: CifContainer):
    """
    This Method adds the descriptions to the fist table column.
    """
    lgnd1 = table.cell(1, 0).paragraphs[0]
    lgnd1sub = lgnd1.add_run('Empirical formula')
    lgnd2 = table.cell(2, 0).paragraphs[0]
    lgnd2sub = lgnd2.add_run('Formula weight')
    lgnd3 = table.cell(3, 0).paragraphs[0]
    lgnd3sub = lgnd3.add_run('Temperature/K')
    lgnd4 = table.cell(4, 0).paragraphs[0]
    lgnd4sub = lgnd4.add_run('Crystal system')
    lgnd5 = table.cell(5, 0).paragraphs[0]
    lgnd5sub = lgnd5.add_run('Space group (number)')
    lgnd6 = table.cell(6, 0).paragraphs[0]
    lgnd6sub = lgnd6.add_run('a').font.italic = True
    lgnd6.add_run('/Å')
    lgnd7 = table.cell(7, 0).paragraphs[0]
    lgnd7sub = lgnd7.add_run('b').font.italic = True
    lgnd7.add_run('/Å')
    lgnd8 = table.cell(8, 0).paragraphs[0]
    lgnd8sub = lgnd8.add_run('c').font.italic = True
    lgnd8.add_run('/Å')
    lgnd9 = table.cell(9, 0).paragraphs[0]
    lgnd9sub = lgnd9.add_run('\u03B1/Å')
    lgnd10 = table.cell(10, 0).paragraphs[0]
    lgnd10sub = lgnd10.add_run('\u03B2/Å')
    lgnd11 = table.cell(11, 0).paragraphs[0]
    lgnd11sub = lgnd11.add_run('\u03B3/Å')
    lgnd12 = table.cell(12, 0).paragraphs[0]
    lgnd12sub = lgnd12.add_run('Volume/Å')
    lgnd12sub1 = lgnd12.add_run('3')
    lgnd12sub1.font.superscript = True
    lgnd13 = table.cell(13, 0).paragraphs[0]
    lgnd13sub = lgnd13.add_run('Z')
    lgnd13sub.font.italic = True
    lgnd14 = table.cell(14, 0).paragraphs[0]
    lgnd14sub = lgnd14.add_run('\u03C1')
    lgnd14sub.font.italic = True
    lgnd14sub1 = lgnd14.add_run('calc')
    lgnd14sub1.font.subscript = True
    lgnd14sub2 = lgnd14.add_run(' g/cm')
    lgnd14sub3 = lgnd14.add_run('3')
    lgnd14sub3.font.superscript = True
    lgnd15 = table.cell(15, 0).paragraphs[0]
    lgnd15sub = lgnd15.add_run('\u03BC/mm')
    lgnd15sub1 = lgnd15.add_run('-1')
    lgnd15sub1.font.superscript = True
    lgnd16 = table.cell(16, 0).paragraphs[0]
    lgnd16sub = lgnd16.add_run('F')
    lgnd16sub.font.italic = True
    lgnd16sub1 = lgnd16.add_run('(000)')
    lgnd17 = table.cell(17, 0).paragraphs[0]
    lgnd17sub = lgnd17.add_run('Crystal size/mm')
    lgnd17sub1 = lgnd17.add_run('3')
    lgnd17sub1.font.superscript = True
    lgnd18 = table.cell(18, 0).paragraphs[0]
    lgnd18sub = lgnd18.add_run('Crystal colour')
    lgnd19 = table.cell(19, 0).paragraphs[0]
    lgnd19sub = lgnd19.add_run('Crystal shape')
    lgnd20 = table.cell(20, 0).paragraphs[0]
    lgnd20sub = lgnd20.add_run('Radiation')
    lgnd21 = table.cell(21, 0).paragraphs[0]
    lgnd21sub = lgnd21.add_run('2\u03F4 range/\u00b0')
    lgnd22 = table.cell(22, 0).paragraphs[0]
    lgnd22sub = lgnd22.add_run('Index ranges')
    lgnd23 = table.cell(23, 0).paragraphs[0]
    lgnd23sub = lgnd23.add_run('Reflections collected')
    lgnd24 = table.cell(24, 0).paragraphs[0]
    lgnd24sub = lgnd24.add_run('Independent reflections')
    lgnd25 = table.cell(25, 0).paragraphs[0]
    lgnd25sub = lgnd25.add_run('Data / Restraints / Param.')
    lgnd26 = table.cell(26, 0).paragraphs[0]
    lgnd26sub = lgnd26.add_run('Goodness-of-fit on ')
    lgnd26sub1 = lgnd26.add_run('F')
    lgnd26sub1.font.italic = True
    lgnd26sub2 = lgnd26.add_run('2')
    lgnd26sub2.font.superscript = True
    lgnd27 = table.cell(27, 0).paragraphs[0]
    lgnd27sub = lgnd27.add_run('Final ')
    lgnd27sub1 = lgnd27.add_run('R')
    lgnd27sub1.font.italic = True
    lgnd27sub2 = lgnd27.add_run(' indexes \n[')
    lgnd27sub3 = lgnd27.add_run('I')
    lgnd27sub3.font.italic = True
    lgnd27sub4 = lgnd27.add_run('\u22652\u03C3(')
    lgnd27sub5 = lgnd27.add_run('I')
    lgnd27sub5.font.italic = True
    lgnd27sub3 = lgnd27.add_run(')]')
    lgnd28 = table.cell(28, 0).paragraphs[0]
    lgnd28sub = lgnd28.add_run('Final ')
    lgnd28sub1 = lgnd28.add_run('R')
    lgnd28sub1.font.italic = True
    lgnd28sub2 = lgnd28.add_run(' indexes \n[all data]')
    lgnd29 = table.cell(29, 0).paragraphs[0]
    lgnd29sub = lgnd29.add_run('Largest peak/hole /eÅ')
    lgnd29sub1 = lgnd29.add_run('3')
    lgnd29sub1.font.superscript = True
    lgnd30 = table.cell(30, 0).paragraphs[0]
    if not cif.is_centrosymm:
        lgnd30sub = lgnd30.add_run('Flack x parameter')


def format_space_group(table, cif):
    """
    Sets formating of the space group symbol in row 6.
    """
    # The HM space group symbol
    space_group = cif['_space_group_name_H-M_alt'].strip("'")
    it_number = cif['_space_group_IT_number']
    if space_group:
        if len(space_group) > 4:  # don't modify P 1
            space_group = re.sub(r'\s1', '', space_group)  # remove extra Hall "1" for mono and tric
        space_group = re.sub(r'\s', '', space_group)  # remove all remaining whitespace
        # space_group = re.sub(r'-1', u'\u0031\u0305', space_group)  # exchange -1 with 1bar
        space_group_formated_text = [char for char in space_group]  # ???)
        sgrun = table.cell(5, 1).paragraphs[0]
        is_sub = False
        for k, char in enumerate(space_group_formated_text):
            sgrunsub = sgrun.add_run(char)
            if not char.isdigit():
                sgrunsub.font.italic = True
            else:
                if space_group_formated_text[k - 1].isdigit() and not is_sub:
                    is_sub = True
                    sgrunsub.font.subscript = True  # lowercase the second digit if previous is also digit
                else:
                    is_sub = False  # only every second number as subscript for P212121 etc.
        if it_number:
            sgrun.add_run(' (' + it_number + ')')
    else:
        space_group = 'no space group'
    return space_group


def make_report_from(file_obj: Path, output_filename: str = None):
    """
    Creates a tabular cif report.
    :param file_obj: Input cif file.
    :param output_filename: the table is saved to this file.
    """
    document = Document('template/template1.docx')
    # Deleting first (empty) paragraph, otherwise first line would be an empty one:
    p = document.paragraphs[0]
    delete_paragraph(p)
    style = document.styles['Normal']
    font = style.font
    #font.name = 'Arial'
    #font.size = Pt(10)
    document.add_heading('Structure Tables', 1)

    # a style for the header:
    styles = document.styles
    # new_heading_style = styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
    # new_heading_style.base_style = styles['Heading 1']
    # font = new_heading_style.font
    # font.color.rgb = RGBColor(0, 0, 0)

    cif = None
    if file_obj.exists():
        try:
            cif = CifContainer(file_obj)
        except Exception as e:
            print('Unable to open cif file:')
            print(e)
    else:
        raise FileNotFoundError

    table_num = 1
    cif, table_num = add_main_table(document, cif, table_num)

    table_num = add_coords_table(document, cif, table_num)

    table_num = add_bonds_and_angles_table(document, cif, table_num)

    table_num = add_torsion_angles(document, cif, table_num)

    table_num = add_hydrogen_bonds(document, cif, table_num)

    print('\nScript finished - output file: tables.docx')
    if not output_filename:
        document.save('tables.docx')
    else:
        document.save(output_filename)
    return file_obj.name

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_main_table(document: Document(), cif: CifContainer, table_num: int):
    tab0_head = r"Table {}. Crystal data and structure refinement for {}".format(table_num, cif.fileobj.name)
    document.add_heading(tab0_head, 2)
    table = document.add_table(rows=31, cols=2)
    #table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.width = Cm(10)
    #table.allow_autofit = False
    #table.autofit = False
    #col = table.columns[0]
    #col.width = Cm(5.0)
    #col.autofit = False
    #col = table.columns[1]
    #col.width = Cm(5.0)
    #col.autofit = False
    # setup table format:
    widths = (Cm(5), Cm(5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    # Add descriptions to the first column of the main table:
    populate_description_columns(table, cif)
    # The main residuals table:
    populate_main_table_values(table, cif)
    return cif, table_num


def populate_main_table_values(table: table, cif: CifContainer):
    """
    Fills the main table with residuals. Column, by column.
    """
    header_cells = table.rows[0].cells
    headcell = header_cells[1].paragraphs[0]
    headcell.add_run(cif.fileobj.name).bold = True

    # Set text for all usual cif keywords by a lookup table:
    for _, key in enumerate(cif_keywords_list):
        # key[1] contains the row number:
        cell = table.cell(key[1] + 1, 1)
        if cif[key[0]]:
            cell.text = cif[key[0]]
        else:
            cell.text = '?'
            continue
    # Now the special handling:
    # The sum formula:
    sum_formula = 'no sum formula'
    if cif['_chemical_formula_sum']:
        sum_formula = cif['_chemical_formula_sum']
        ltext2 = sum_formula.replace(" ", "").replace("'", "")
        ltext3 = [''.join(x[1]) for x in it.groupby(ltext2, lambda x: x.isalpha())]
        for _, word in enumerate(ltext3):
            formrun = table.cell(1, 1).paragraphs[0]
            formrunsub = formrun.add_run(word)
            if isfloat(word):
                formrunsub.font.subscript = True

    space_group = format_space_group(table, cif)
    radiation_type = cif['_diffrn_radiation_type']
    radiation_wavelength = cif['_diffrn_radiation_wavelength']
    crystal_size_min = cif['_exptl_crystal_size_min']
    crystal_size_mid = cif['_exptl_crystal_size_mid']
    crystal_size_max = cif['_exptl_crystal_size_max']
    limit_h_min = cif['_diffrn_reflns_limit_h_min']
    limit_h_max = cif['_diffrn_reflns_limit_h_max']
    limit_k_min = cif['_diffrn_reflns_limit_k_min']
    limit_k_max = cif['_diffrn_reflns_limit_k_max']
    theta_min = cif['_diffrn_reflns_theta_min']
    theta_max = cif['_diffrn_reflns_theta_max']
    limit_l_min = cif['_diffrn_reflns_limit_l_min']
    limit_l_max = cif['_diffrn_reflns_limit_l_max']
    reflns_number_total = cif['_reflns_number_total']
    reflns_av_R_equivalents = cif['_diffrn_reflns_av_R_equivalents']
    reflns_av_unetI = cif['_diffrn_reflns_av_unetI/netI']
    ls_number_reflns = cif['_refine_ls_number_reflns']
    ls_number_restraints = cif['_refine_ls_number_restraints']
    ls_number_parameters = cif['_refine_ls_number_parameters']
    ls_R_factor_gt = cif['_refine_ls_R_factor_gt']
    ls_wR_factor_gt = cif['_refine_ls_wR_factor_gt']
    ls_R_factor_all = cif['_refine_ls_R_factor_all']
    ls_wR_factor_ref = cif['_refine_ls_wR_factor_ref']
    try:
        diff_density_min = "{0:.2f}".format(round(float(cif['_refine_diff_density_min']), 2))
    except ValueError:
        diff_density_min = '?'
    try:
        diff_density_max = "{0:.2f}".format(round(float(cif['_refine_diff_density_max']), 2))
    except ValueError:
        diff_density_max = '?'

    # now prepare & write all the concatenated & derived cell contents:
    table.cell(17, 1).text = this_or_quest(crystal_size_max) + '\u00d7' + \
                             this_or_quest(crystal_size_mid) + '\u00d7' + \
                             this_or_quest(crystal_size_min)
    wavelength = str(' (\u03bb =' + this_or_quest(radiation_wavelength) + ')').replace(' ', '')
    # radtype: ('Mo', 'K', '\\a')
    radtype = list(radiation_type.partition("K"))
    if len(radtype) > 2:
        if radtype[2] == '\\a':
            radtype[2] = '\u03b1'
        if radtype[2] == '\\b':
            radtype[2] = '\u03b2'
    radrun = table.cell(20, 1).paragraphs[0]
    # radiation type e.g. Mo:
    radrun.add_run(radtype[0])
    # K line:
    radrunita = radrun.add_run(radtype[1])
    radrunita.font.italic = True
    alpha = radrun.add_run(radtype[2])
    alpha.font.italic = True
    alpha.font.subscript = True
    # wavelength lambda:
    radrun.add_run(' ' + wavelength)
    try:
        table.cell(21, 1).text = "{0:.2f}".format(2 * float(theta_min)) + \
                                 ' to ' + "{0:.2f}".format(2 * float(theta_max))
    except ValueError:
        table.cell(21, 1).text = '? to ?'
    table.cell(22, 1).text = limit_h_min + ' \u2264 h \u2264 ' \
                             + limit_h_max + '\n' \
                             + limit_k_min + ' \u2264 k \u2264 ' \
                             + limit_k_max + '\n' \
                             + limit_l_min + ' \u2264 l \u2264 ' \
                             + limit_l_max
    rintrun = table.cell(24, 1).paragraphs[0]
    rintrun.add_run(this_or_quest(reflns_number_total) + '\n')
    rintita1 = rintrun.add_run('R')
    rintita1.font.italic = True
    rintsub1 = rintrun.add_run('int')
    rintsub1.font.subscript = True
    rintrun.add_run(' = ' + this_or_quest(reflns_av_R_equivalents) + '\n')
    rintita2 = rintrun.add_run('R')
    rintita2.font.italic = True
    rintsub2 = rintrun.add_run('sigma')
    rintsub2.font.subscript = True
    rintrun.add_run(' = ' + this_or_quest(reflns_av_unetI))
    table.cell(25, 1).text = this_or_quest(ls_number_reflns) + '/' \
                             + this_or_quest(ls_number_restraints) + '/' \
                             + this_or_quest(ls_number_parameters)
    r2sigrun = table.cell(27, 1).paragraphs[0]
    r2sigita1 = r2sigrun.add_run('R')
    r2sigita1.font.italic = True
    r2sigsub1 = r2sigrun.add_run('1')
    r2sigsub1.font.subscript = True
    r2sigrun.add_run(' = ' + this_or_quest(ls_R_factor_gt) + '\nw')
    r2sigita2 = r2sigrun.add_run('R')
    r2sigita2.font.italic = True
    r2sigsub2 = r2sigrun.add_run('2')
    r2sigsub2.font.subscript = True
    r2sigrun.add_run(' = ' + this_or_quest(ls_wR_factor_gt))
    rfullrun = table.cell(28, 1).paragraphs[0]
    rfullita1 = rfullrun.add_run('R')
    rfullita1.font.italic = True
    rfullsub1 = rfullrun.add_run('1')
    rfullsub1.font.subscript = True
    rfullrun.add_run(' = ' + this_or_quest(ls_R_factor_all) + '\nw')
    rfullita2 = rfullrun.add_run('R')
    rfullita2.font.italic = True
    rfullsub2 = rfullrun.add_run('2')
    rfullsub2.font.subscript = True
    rfullrun.add_run(' = ' + ls_wR_factor_ref)
    table.cell(29, 1).text = diff_density_max + '/' + diff_density_min
    if not cif.is_centrosymm:
        table.cell(30, 1).text = cif['_refine_ls_abs_structure_Flack'] or '?'
    print('File parsed: ' + cif.fileobj.name + '  (' + sum_formula + ')  ' + space_group)


def add_coords_table(document: Document, cif: CifContainer, table_num: int):
    """
    Adds the table with the atom coordinates.
    :param document: The current word document.
    :param cif: the cif object from CifContainer.
    :return: None
    """
    table_num += 1
    # TODO: proper formating for heading:
    headline = "Table {}. Atomic coordinates ".format(table_num)
    h = document.add_heading(headline, 2)
    h.add_run(' and equivalent isotropic displacement parameters (Å')
    h.add_run('2').font.superscript = True
    h.add_run(') for {}. U'.format(cif.fileobj.name))
    h.add_run('eq').font.subscript = True
    h.add_run(' is defined as one third of the trace of the orthogonalized U')
    ij = h.add_run('ij')
    ij.font.subscript = True
    ij.italic = True
    h.add_run(' tensor.')
    coords_table = document.add_table(rows=1, cols=5)
    #coords_table.style = document.styles['Table Grid']
    coords_table.style = 'Table Grid'
    # Atom	x	y	z	U(eq)
    head_row = coords_table.rows[0]
    ar = head_row.cells[0].paragraphs[0].add_run('Atom')
    ar.bold = True
    ar = head_row.cells[1].paragraphs[0].add_run('x')
    ar.bold = True
    ar.italic = True
    ar = head_row.cells[2].paragraphs[0].add_run('y')
    ar.bold = True
    ar.italic = True
    ar = head_row.cells[3].paragraphs[0].add_run('z')
    ar.italic = True
    ar = head_row.cells[4].paragraphs[0].add_run('U')
    ar.bold = True
    ar.italic = True
    ar = head_row.cells[4].paragraphs[0].add_run('eq')
    ar.bold = True
    ar.font.subscript = True
    for at in cif.atoms():
        row_cells = coords_table.add_row().cells
        row_cells[0].text = at[0]  # label
        row_cells[1].text = str(at[2])  # x
        row_cells[2].text = str(at[3])  # y
        row_cells[3].text = str(at[4])  # z
        row_cells[4].text = str(at[7])  # ueq
    return table_num


def add_bonds_and_angles_table(document: Document, cif: CifContainer, table_num: int):
    """
    Make table with bonds and angles.
    """
    table_num += 1
    headline = r"Table {}. Bond lengths [Å] and angles [°] for {}.".format(table_num, cif.fileobj.name)
    document.add_heading(headline, 2)
    bond_angle_table = document.add_table(rows=0, cols=2)
    bond_angle_table.style = 'Table Grid'
    # Bond/Angle  value
    head_row = bond_angle_table.add_row()
    ar = head_row.cells[0].paragraphs[0].add_run('Atom - Atom')
    ar.bold = True
    ar = head_row.cells[1].paragraphs[0].add_run('Length [Å]')
    ar.bold = True
    symms = {}
    newsymms = {}
    num = 1
    for at1, at2, val, symm2 in cif.bonds():
        if symm2 == '.':
            symm2 = None
        if symm2 and symm2 not in symms.keys():
            symms[symm2] = num
            # Applys translational symmetry to symmcards:
            # 3_556 -> 2
            card = get_card(cif, symm2)
            s = SymmetryElement(card)
            s.translate(symm2)
            newsymms[num] = s.toShelxl()
            num += 1
        row_cells = bond_angle_table.add_row().cells
        row_cells[0].text = at1 + ' - ' + at2
        row_cells[0].paragraphs[0].add_run('#' + str(symms[symm2]) if symm2 else '').font.superscript = True
        row_cells[1].text = str(val)  # bond
    ############ the angles ####################
    bond_angle_table.add_row()
    head_row = bond_angle_table.add_row()
    ar = head_row.cells[0].paragraphs[0].add_run('Atom - Atom - Atom')
    ar.bold = True
    ar = head_row.cells[1].paragraphs[0].add_run('Angle [°]')
    ar.bold = True
    card = ''
    s = SymmetryElement(card)
    # TODO: split this in two columns:
    for at1, at2, at3, angle, symm1, symm3 in cif.angles():
        if symm1 == '.':
            symm1 = None
        if symm3 == '.':
            symm3 = None
        if (symm1 or symm3) and (symm1 or symm3) not in symms.keys():
            if symm1:
                symms[symm1] = num
                card = get_card(cif, symm1)
            if symm3:
                symms[symm3] = num
                card = get_card(cif, symm3)
            # Applys translational symmetry to symmcards:
            # 3_556 -> 2
            if symm1 or symm3:
                s = SymmetryElement(card)
            if symm1:
                s.translate(symm1)
            if symm3:
                s.translate(symm3)
            newsymms[num] = s.toShelxl()
            num += 1
        row_cells = bond_angle_table.add_row().cells
        row_cells[0].text = at1
        row_cells[0].paragraphs[0].add_run('#' + str(symms[symm1]) if symm1 else '').font.superscript = True
        row_cells[0].paragraphs[0].add_run(' - ' + at2 + ' - ')
        row_cells[0].paragraphs[0].add_run(at3)  # labels
        row_cells[0].paragraphs[0].add_run('#' + str(symms[symm3]) if symm3 else '').font.superscript = True
        row_cells[1].text = str(angle)  # angle
    widths = (Cm(4), Cm(4))
    make_table_widths(bond_angle_table, widths)
    add_last_symminfo_line(newsymms, document)
    return table_num


def make_table_widths(table, widths):
    """
    Sets the width of the columns of a table.
    """
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_last_symminfo_line(newsymms, document):
    p = document.add_paragraph('')
    line = 'Symmetry transformations used to generate equivalent atoms: '
    nitems = len(newsymms)
    n = 0
    for key, value in newsymms.items():
        sep = ';'
        if n == nitems:
            sep = ''
        n += 1
        line += "#{}: {}{}   ".format(key, value, sep)
    if newsymms:
        p.add_run(line)


def get_card(cif, symm):
    card = cif.symmops[int(symm.split('_')[0]) - 1].split(',')
    return card


def add_torsion_angles(document: Document, cif: CifContainer, table_num: int):
    """
    Table 6.  Torsion angles [°] for I-43d_final.
    """
    if not len(list(cif.torsion_angles())) > 0:
        print('No torsion angles in cif.')
        return table_num
    table_num += 1
    headline = r"Table {}. Torsion angles [°] for {}.".format(table_num, cif.fileobj.name)
    document.add_heading(headline, 2)
    torsion_table = document.add_table(rows=0, cols=2)
    symms = {}
    newsymms = {}
    card = ''
    s = SymmetryElement(card)
    num = 1
    for at1, at2, at3, at4, angle, symm1, symm2, symm3, symm4 in cif.torsion_angles():
        if symm1 == '.':
            symm1 = None
        if symm2 == '.':
            symm2 = None
        if symm3 == '.':
            symm3 = None
        if symm4 == '.':
            symm4 = None
        if (symm1 or symm2 or symm3 or symm4) and (symm1 or symm2 or symm3 or symm4) not in symms.keys():
            if symm1:
                symms[symm1] = num
                s = SymmetryElement(get_card(cif, symm1))
                s.translate(symm1)
            if symm2:
                symms[symm2] = num
                s = SymmetryElement(get_card(cif, symm2))
                s.translate(symm2)
            if symm3:
                symms[symm3] = num
                s = SymmetryElement(get_card(cif, symm3))
                s.translate(symm3)
            if symm4:
                symms[symm4] = num
                s = SymmetryElement(get_card(cif, symm4))
                s.translate(symm4)
            newsymms[num] = s.toShelxl()
            num += 1
        row_cells = torsion_table.add_row().cells
        row_cells[0].text = at1
        row_cells[0].paragraphs[0].add_run('#' + str(symms[symm1]) if symm1 else '').font.superscript = True
        row_cells[0].paragraphs[0].add_run(' - ')
        row_cells[0].paragraphs[0].add_run(at2)
        row_cells[0].paragraphs[0].add_run('#' + str(symms[symm2]) if symm2 else '').font.superscript = True
        row_cells[0].paragraphs[0].add_run(' - ')
        row_cells[0].paragraphs[0].add_run(at3)
        row_cells[0].paragraphs[0].add_run('#' + str(symms[symm3]) if symm3 else '').font.superscript = True
        row_cells[0].paragraphs[0].add_run(' - ')
        row_cells[0].paragraphs[0].add_run(at4)  # labels
        row_cells[0].paragraphs[0].add_run('#' + str(symms[symm4]) if symm4 else '').font.superscript = True
        row_cells[1].paragraphs[0].add_run(str(angle))  # angle
    widths = (Cm(4), Cm(4))
    make_table_widths(torsion_table, widths)
    add_last_symminfo_line(newsymms, document)
    return table_num


def add_hydrogen_bonds(document: Document, cif: CifContainer, table_num: int):
    """
    Table 7.  Hydrogen bonds for I-43d_final  [Å and °].
    """
    if not len(list(cif.hydrogen_bonds())) > 0:
        print('No hydrogen bonds in cif.')
        return
    table_num += 1
    headline = r"Table {}. Hydrogen bonds for {} [Å and °].".format(table_num, cif.fileobj.name)
    document.add_heading(headline, 2)
    hydrogen_table = document.add_table(rows=1, cols=5)
    head_row = hydrogen_table.rows[0].cells
    # D-H...A	d(D-H)	d(H...A)	d(D...A)	<(DHA)
    head_row[0].paragraphs[0].add_run('D-H...A').font.bold = True
    head_row[1].paragraphs[0].add_run('d(D-H)').font.bold = True
    head_row[2].paragraphs[0].add_run('d(H...A)').font.bold = True
    head_row[3].paragraphs[0].add_run('d(D...A)').font.bold = True
    head_row[4].paragraphs[0].add_run('<(DHA)').font.bold = True
    symms = {}
    newsymms = {}
    num = 1
    for label_d, label_h, label_a, dist_dh, dist_ha, dist_da, angle_dha, symm in cif.hydrogen_bonds():
        if symm == '.':
            symm = None
        if symm and symm not in symms.keys():
            symms[symm] = num
            s = SymmetryElement(get_card(cif, symm))
            s.translate(symm)
            newsymms[num] = s.toShelxl()
        num += 1
        symmval = ('#' + str(symms[symm])) if symm else ''
        row_cells = hydrogen_table.add_row().cells
        row_cells[0].text = label_d + ' - ' + label_h + ' ... ' + label_a
        row_cells[0].paragraphs[0].add_run(symmval).font.superscript = True
        row_cells[1].text = dist_dh
        row_cells[2].text = dist_ha
        row_cells[3].text = dist_da
        row_cells[4].text = angle_dha
    add_last_symminfo_line(newsymms, document)
    return table_num


if __name__ == '__main__':
    output_filename = 'tables.docx'
    # make_report_from(get_files_from_current_dir()[1])
    make_report_from(Path(r'test-data/DK_zucker2_0m.cif'))
    # make_report_from(Path(r'/Volumes/home/strukturen/eigene/DK_30011/sad-final.cif'))
    # make_report_from(Path(r'D:\goedaten\strukturen_goe\eigene\DK_4008\xl12\new\r3c.cif'))
    if sys.platform == 'win':
        os.startfile(output_filename)
    if sys.platform == 'darwin':
        subprocess.call(['open', output_filename])
