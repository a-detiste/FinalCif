from builtins import str

import gemmi
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from cif.file_reader import CifContainer

"""
TODO: Add references of the used programs to the end.
"""


class FormatMixin():

    def bold(self, run: Run):
        r = run.bold = True
        return r


class ReportText():
    def __init__(self):
        pass


class CrstalSelection(FormatMixin):
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        self.temp = self.cif['_diffrn_ambient_temperature']
        self._name = cif.fileobj.name
        method = 'shock-cooled '
        sentence = "The data for {} were collected from a {}single crystal at {}\u00A0K."
        if float(self.temp.split('(')[0]) > 200:
            method = ''
        self.txt = sentence.format(self.name, method, self.temp)
        paragraph.add_run(self.txt)

    @property
    def name(self) -> str:
        if self._name:
            return self._name
        else:
            return '?'


class MachineType():
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        gstr = gemmi.cif.as_string
        self.difftype = gstr(self.cif.block.find_value('_diffrn_measurement_device_type'))
        self.device = gstr(self.cif['_diffrn_measurement_device'])
        self.source = gstr(self.cif['_diffrn_source'])
        self.monochrom = gstr(self.cif['_diffrn_radiation_monochromator'])
        if not self.monochrom:
            self.monochrom = '?'
        self.cooling = gstr(self.cif['_olex2_diffrn_ambient_temperature_device'])
        if not self.cooling:
            self.cooling = '?'
        self.rad_type = gstr(self.cif['_diffrn_radiation_type'])
        radtype = format_radiation(self.rad_type)
        self.wavelen = gstr(self.cif['_diffrn_radiation_wavelength'])
        sentence1 = "The data were collected on {} {} {} with {} {} using {} as monochromator. " \
                    "The diffractometer were equipped with {} {} low temperature device and used "
        sentence2 = " radiation, λ = {}\u00A0Å. "
        txt = sentence1.format(get_inf_article(self.difftype), self.difftype, self.device,
                               get_inf_article(self.source), self.source, self.monochrom,
                               get_inf_article(self.cooling), self.cooling)
        paragraph.add_run(txt)
        # radiation type e.g. Mo:
        paragraph.add_run(radtype[0])
        # K line:
        radrunita = paragraph.add_run(radtype[1])
        radrunita.font.italic = True
        alpha = paragraph.add_run(radtype[2])
        alpha.font.italic = True
        alpha.font.subscript = True
        txt2 = sentence2.format(self.wavelen)
        paragraph.add_run(txt2)


class DataReduct():
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        gstr = gemmi.cif.as_string
        integration = gstr(self.cif.block.find_value('_computing_data_reduction'))
        integration_prog = '?'
        if 'saint' in integration.lower():
            integration_prog = 'SAINT'
        if 'crysalis' in integration.lower():
            integration_prog = 'CrysAlisPro'
        if 'trek' in integration.lower():
            integration_prog = 'd*trek'
        abstype = gstr(self.cif.block.find_value('_exptl_absorpt_correction_type'))
        abs_details = gstr(self.cif.block.find_value('_exptl_absorpt_process_details'))
        if 'sortav' in abs_details.lower():
            abs_details = 'SORTAV'
        if 'sadabs' in abs_details.lower():
            abs_details = 'SADABS'
        if 'twinabs' in abs_details.lower():
            abs_details = 'TWINABS'
        if 'crysalis' in abs_details.lower():
            abs_details = 'SCALE3 ABSPACK'
        sentence = 'All data were integrated with {} and {} {} absorption correction using {} was applied. '
        txt = sentence.format(integration_prog, get_inf_article(abstype), abstype, abs_details)
        paragraph.add_run(txt)


class SolveRefine():
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        gstr = gemmi.cif.as_string
        solution_prog = gstr(self.cif.block.find_value('_computing_structure_solution'))
        if not solution_prog:
            solution_prog = '?'
        solution_method = gstr(self.cif.block.find_value('_atom_sites_solution_primary'))
        if not solution_method:
            solution_method = '?'
        refined = gstr(self.cif.block.find_value('_computing_structure_refinement'))
        if not refined:
            refined = '?'
        # dsr = gstr(self.cif.block.find_value('_computing_structure_refinement'))
        refine_coef = gstr(self.cif.block.find_value('_refine_ls_structure_factor_coef'))
        sentence = r"The structure were solved by {} methods using {} and refined by full-matrix " \
                   "least-squares methods against "
        txt = sentence.format(solution_method, solution_prog)
        paragraph.add_run(txt)
        paragraph.add_run('F').font.italic = True
        if refine_coef.lower() == 'fsqd':
            paragraph.add_run('2').font.superscript = True
        paragraph.add_run(' by {}'.format(refined))
        paragraph.add_run('. ')


class Hydrogens():
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        """
        TODO: check if the proposed things are really there.
        """
        self.cif = cif
        sentence1 = "All non-hydrogen atoms were refined with anisotropic displacement parameters. " \
                    "The hydrogen atoms were refined isotropically on calculated positions using a riding model " \
                    "with their "  # iso values constrained to 1.5 times the Ueq of their pivot atoms for " \
        # "terminal sp3 carbon atoms and 1.2 times for all other carbon atoms."
        sentence2 = " values constrained to 1.5 times the "
        sentence3 = " of their pivot atoms for terminal sp"
        sentence4 = " carbon atoms and 1.2 times for all other carbon atoms."
        paragraph.add_run(sentence1)
        paragraph.add_run('U').font.italic = True
        paragraph.add_run('iso').font.subscript = True
        paragraph.add_run(sentence2)
        paragraph.add_run('U').font.italic = True
        paragraph.add_run('eq').font.subscript = True
        paragraph.add_run(sentence3)
        paragraph.add_run('3').font.superscript = True
        paragraph.add_run(sentence4)


def get_inf_article(next_word: str) -> str:
    if not next_word:
        return 'a'
    voc = 'aeiou'
    return 'a' if not next_word[0].lower() in voc else 'an'


def format_radiation(radiation_type: str) -> list:
    radtype = list(radiation_type.partition("K"))
    if len(radtype) > 2:
        if radtype[2] == r'\a':
            radtype[2] = '\u03b1'
        if radtype[2] == r'\b':
            radtype[2] = '\u03b2'
    return radtype
