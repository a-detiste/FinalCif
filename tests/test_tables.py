import unittest
from pathlib import Path

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shape import InlineShapes
from docx.shared import Cm
from docx.table import Table

from finalcif import VERSION
from finalcif.appwindow import AppWindow


class TablesTestMixin():

    def setUp(self) -> None:
        self.testcif = Path('tests/examples/1979688.cif').absolute()
        self.myapp = AppWindow(self.testcif, unit_test=True)
        self.myapp.ui.HAtomsCheckBox.setChecked(False)
        self.myapp.ui.ReportTextCheckBox.setChecked(False)
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(0.0)
        # make sure to use no template:
        self.myapp.ui.TemplatesListWidget.setCurrentRow(0)
        self.myapp.running_inside_unit_test = True
        self.myapp.hide()
        self.reportdoc = self.myapp.cif.finalcif_file_prefixed(prefix='report_', suffix='-finalcif.docx')
        self.report_zip = self.myapp.cif.finalcif_file_prefixed(prefix='', suffix='-finalcif.zip')
        self.myapp.hide()

    def tearDown(self) -> None:
        self.myapp.cif.finalcif_file.unlink(missing_ok=True)
        self.reportdoc.unlink(missing_ok=True)
        self.report_zip.unlink(missing_ok=True)
        self.myapp.ui.ReportTextCheckBox.setChecked(False)
        self.myapp.ui.HAtomsCheckBox.setChecked(False)
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(7.5)
        self.myapp.close()


class TablesTestCase(TablesTestMixin, unittest.TestCase):

    def setUp(self) -> None:
        super().setUp()
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(7.43)
        self.myapp.select_report_picture(Path('finalcif/icon/finalcif.png'))

    def test_save_report_works(self):
        self.myapp.ui.SaveFullReportButton.click()
        self.assertEqual(True, self.reportdoc.exists())

    def test_picture_has_correct_size(self):
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(WD_INLINE_SHAPE.PICTURE, shapes[0].type)
        self.assertEqual(Cm(7.43).emu, shapes[0].width)

    def test_default_picture_width(self):
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(0.0)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.resolve())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(WD_INLINE_SHAPE.PICTURE, shapes[0].type)
        self.assertEqual(Cm(7.5).emu, shapes[0].width)

    def test_option_with_h(self):
        self.myapp.ui.HAtomsCheckBox.setChecked(False)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        table: Table = doc.tables[3]
        self.assertEqual('C1–H1', table.cell(row_idx=4, col_idx=0).text)

    def test_option_without_h(self):
        self.myapp.ui.HAtomsCheckBox.setChecked(True)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        table: Table = doc.tables[3]
        self.assertEqual('O1–C13', table.cell(row_idx=4, col_idx=0).text)

    def test_option_no_report_text(self):
        self.myapp.ui.ReportTextCheckBox.setChecked(True)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        self.assertEqual('The methanol molecule is d', doc.paragraphs[3].text[:26])

    def test_option_no_report_text_ueq(self):
        self.myapp.ui.ReportTextCheckBox.setChecked(True)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        self.assertEqual('Ueq is defined as 1/3 of t', doc.paragraphs[5].text[:26])

    def test_option_with_report_text(self):
        self.myapp.ui.ReportTextCheckBox.setChecked(False)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        self.assertEqual('The following text is only', doc.paragraphs[3].text[:26])

    def test_citations(self):
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        self.assertEqual(
            '[6] 	D. Kratzert, FinalCif, V{}, https://dkratzert.de/finalcif.html.'.format(VERSION),
            doc.paragraphs[-1].text)

    def test_ccdc_num_in_table(self):
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        table: Table = doc.tables[0]
        self.assertEqual('1979688', table.cell(row_idx=0, col_idx=1).text)


class TablesNoPictureTestCase(TablesTestMixin, unittest.TestCase):

    def setUp(self) -> None:
        super().setUp()

    def test_save_report_works(self):
        self.myapp.ui.SaveFullReportButton.click()
        self.assertEqual(True, self.reportdoc.exists())

    def test_picture_has_correct_size(self):
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.resolve())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(0, len(shapes))


if __name__ == '__main__':
    unittest.main()
